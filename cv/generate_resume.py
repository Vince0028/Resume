from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Page margins (tight for 1-page fit) ──────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin   = Inches(0.85)
    section.right_margin  = Inches(0.85)

# ── Default paragraph style ───────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after  = Pt(0)
normal.paragraph_format.space_before = Pt(0)

CRIMSON  = RGBColor(0x8B, 0x00, 0x00)
TAB_POS  = '8640'   # ~6 in from left edge at 0.85 margin

# ── Helpers ───────────────────────────────────────────────────────────────────
def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)

def section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text.upper())
    run.bold            = True
    run.font.size       = Pt(10.5)
    run.font.color.rgb  = RGBColor(0x00, 0x00, 0x00)
    run.font.name       = 'Times New Roman'
    add_hr(doc)

def two_col(doc, left, right, left_bold=True, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(0)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab  = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), TAB_POS)
    tabs.append(tab)
    pPr.append(tabs)
    rl = p.add_run(left)
    rl.bold = left_bold; rl.font.name = 'Times New Roman'; rl.font.size = Pt(size)
    p.add_run('\t')
    rr = p.add_run(right)
    rr.font.name = 'Times New Roman'; rr.font.size = Pt(size)
    return p

def sub(doc, text, italic=True, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.italic = italic; r.font.name = 'Times New Roman'; r.font.size = Pt(size)

def bul(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)

def gap(doc, pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(pts)
    p.paragraph_format.space_after  = Pt(0)

def plain(doc, label, body, lsize=10.5, bsize=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(0)
    r1 = p.add_run(label)
    r1.bold = True; r1.font.name = 'Times New Roman'; r1.font.size = Pt(lsize)
    r2 = p.add_run(body)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(bsize)

# ===============================================================
# HEADER
# ===============================================================
hp = doc.add_paragraph()
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
hp.paragraph_format.space_before = Pt(0)
hp.paragraph_format.space_after  = Pt(2)
hr = hp.add_run('VINCE NELMAR ALOBIN')
hr.bold = True; hr.font.size = Pt(15); hr.font.name = 'Times New Roman'
hr.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

cp = doc.add_paragraph()
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp.paragraph_format.space_before = Pt(0)
cp.paragraph_format.space_after  = Pt(1)
cr = cp.add_run('Pasay City, Philippines  |  alobinvince@gmail.com  |  09276922670')
cr.font.size = Pt(9.5); cr.font.name = 'Times New Roman'

add_hr(doc)

# ===============================================================
# SUMMARY
# ===============================================================
section_heading(doc, 'Summary')
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(3)
r = p.add_run(
    'BSIT student and full-stack developer with production experience across web, mobile, '
    'computer vision, and embedded systems. Delivered 17+ end-to-end projects spanning '
    'AI integration, cloud deployment, and hardware prototyping, backed by 40+ '
    'professional certifications in technology, cloud, and cybersecurity.'
)
r.font.name = 'Times New Roman'; r.font.size = Pt(10)

# ===============================================================
# EXPERIENCE
# ===============================================================
section_heading(doc, 'Experience')

two_col(doc, 'Division Science and Technology Fair (DSTF) - Lead Developer', 'Aug 2025 - Oct 2025')
sub(doc, 'Pasay City South High School  |  Pasay City, Philippines')
bul(doc, 'Served as lead developer within a 4-person team, solely building all 4 DSTF projects end-to-end; '
    'compensated as the sole programmer for a combined group of 14-16 people across teams.')
bul(doc, 'Shipped DengueTect (AI/ML dengue risk, React/Flask), AnaLytics (retail analytics, Next.js/Supabase), '
    'Driver Expression Detector (computer vision, Python/Raspberry Pi), and SmartShut (motion-aware lighting, Arduino).')
bul(doc, 'Managed the full development lifecycle across AI-integrated systems, embedded hardware prototypes, '
    'and full-stack web applications, from initial architecture through final deployment.')

# ===============================================================
# PROJECTS
# ===============================================================
section_heading(doc, 'Projects')

bul(doc, 'Engineered GitChat, a fully offline P2P mesh chat and video call app using Bluetooth for '
    'device discovery and WebRTC for streaming; built in Flutter and Dart with Kotlin, requiring zero internet.')
bul(doc, 'Developed Veriface, an automated attendance system powered by real-time facial recognition; '
    'built with Python, OpenCV, and NumPy, eliminating manual rollcall through computer vision.')

# ===============================================================
# SKILLS
# ===============================================================
section_heading(doc, 'Skills')

bul(doc, 'Languages: Python, JavaScript, TypeScript, Dart, Java, C++, Kotlin.')
bul(doc, 'Tools and Frameworks: React, Flutter, Flask, Next.js, NestJS, OpenCV, Supabase, PostgreSQL, '
    'Docker, Git, AWS, Azure, Vercel, Postman, Arduino, Raspberry Pi.')

# ===============================================================
# CERTIFICATIONS
# ===============================================================
section_heading(doc, 'Certifications')

bul(doc, 'Cloud and Security: AWS Academy Cloud Foundations, Microsoft Azure AI Essentials, Career '
    'Essentials in Generative AI (Microsoft/LinkedIn), CISM Cert Prep, CompTIA Network+ N10-009, Ubuntu Linux (Canonical).')
bul(doc, 'Programming and Design: Java Foundations (JetBrains), JavaScript Foundations (Mozilla), '
    'Become a Full-Stack Web Developer, Adobe InDesign 2024 Professional Certificate, '
    'NC2 Animation - National Qualification (TESDA).')

# ===============================================================
# EDUCATION
# ===============================================================
section_heading(doc, 'Education')

two_col(doc, 'Asia Pacific College (APC)  |  Pasay City, Philippines', '2024 - Present')
sub(doc, 'Bachelor of Science in Information Technology', italic=False)
sub(doc, 'Coursework: Web Programming, Mobile Programming, Data Structures, Data Management, Programming Concepts', italic=True)

# ===============================================================
# SAVE
# ===============================================================
out = r'C:\Users\Vince\Documents\Alobin ICT241\website_resume\Vince_Alobin_Harvard_Resume.docx'
doc.save(out)
print(f'Saved: {out}')
